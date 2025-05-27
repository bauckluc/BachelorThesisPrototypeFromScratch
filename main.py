import time
from openai import OpenAI
from apiKey import OPENAI_KEY
from docx import Document

client = OpenAI(api_key=OPENAI_KEY)

assistant_id = "asst_nah37WsEZ9Jvqg94do3K3afo"

# Prepare word-document
document = Document()


# Create thread
thread = client.beta.threads.create()
thread_id = thread.id

# Start chat loop
while True:
    user_input = input("\nDu: ")

    if user_input.lower() in {"exit", "quit", "bye"}:
        print("Chatbot beendet. Bis bald!")
        break

    # Add user message to document
    document.add_paragraph(f"Du: {user_input}")

    # Add new message
    client.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=user_input
    )

    # Start run
    run = client.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id
    )

    # Wait for run to complete
    while True:
        run_status = client.beta.threads.runs.retrieve(
            thread_id=thread_id,
            run_id=run.id
        )
        if run_status.status == "completed":
            break
        time.sleep(1)  # Wait for 1 second before checking again

    # Print assistant's reply
    messages = client.beta.threads.messages.list(thread_id=thread_id)
    assistant_messages = [
        msg for msg in messages.data if msg.role == "assistant"
    ]

    if assistant_messages:
        latest_reply = assistant_messages[0].content[0].text.value
        print(f"\nAssistant: {latest_reply}")
        document.add_paragraph(f"Assistant: {latest_reply}")

    else:
        print("\nAssistant: (keine Antwort erhalten)")

    # Does the user want to save the conversation?
    save_choice = input("\nMöchtest du den bisherigen Verlauf als Word-Datei speichern? (ja/nein): ").lower()
    if save_choice == "ja":
        filename = input("Wie soll die Datei heißen? (ohne .docx): ")
        filepath = f"{filename}.docx"
        document.save(filepath)
        print(f"Datei erfolgreich gespeichert unter: {filepath}")

